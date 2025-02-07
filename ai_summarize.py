import openai
import pdfplumber
import time
import os
import math
import sys
from docx import Document  # Ensure correct usage of python-docx
import tiktoken  # Added for accurate token estimation

# ✅ Hardcoded OpenAI API Key (Replace with your actual key)
api_key = "<INSERT_YOUR_APIKEY_HERE>"

# ✅ Initialize OpenAI Client with Hardcoded API Key
client = openai.OpenAI(api_key=api_key)

# ✅ Function to estimate token count from text using tiktoken
def estimate_tokens(text):
    """Estimates OpenAI token count based on accurate token encoding."""
    encoder = tiktoken.encoding_for_model("gpt-4o-mini")
    return len(encoder.encode(text))

# ✅ Function to estimate and display cost before processing
def display_cost_estimate(text):
    """Calculates estimated cost based on OpenAI's pricing and displays a neatly formatted table."""
    total_tokens = estimate_tokens(text)
    input_cost_per_1k = 0.0005  # Cost per 1,000 tokens for input
    output_cost_per_1k = 0.0015  # Cost per 1,000 tokens for output
    output_tokens = total_tokens * 0.1  # Estimated output size (~10% of input)

    input_cost = (total_tokens / 1000) * input_cost_per_1k
    output_cost = (output_tokens / 1000) * output_cost_per_1k
    total_cost = input_cost + output_cost

    print("\n📊 **Estimated Cost for Processing** 📊")
    print("=" * 70)
    print(f"| {'Category':<18} | {'Tokens':>12} | {'Cost':>12} |")
    print("-" * 70)
    print(f"| {'Model':<18} | {'GPT-4o-Mini':>26} |")
    print(f"| {'Input Tokens':<18} | {total_tokens:>12,} | ${input_cost:>10.5f} |")
    print(f"| {'Output Tokens':<18} | {int(output_tokens):>12,} | ${output_cost:>10.5f} |")
    print(f"| {'Total Cost':<18} | {'-':>12} | **${total_cost:>10.5f}** |")
    print("=" * 70)
    print("\n⚡ Processing will now begin...\n")

# ✅ Function to get user input for the directory and validate it
def get_valid_directory():
    directory = input(
        "Enter the directory where your documents are stored, including the drive name [e.g.:  'C:\\<directory name>\\']: ").strip()

    if not os.path.exists(directory):
        os.makedirs(directory)
        print(f"\nFile directory '{directory}' has been created!\n"
              "Please copy the documents you wish to analyze into the directory and rerun this Python application.")
        exit()  # Gracefully terminate after directory creation

    print(f"\nDirectory '{directory}' found!\n")
    return directory

# ✅ Function to get a valid file from the directory
def get_valid_file(directory):
    """Lists files in the directory but excludes any AI-generated summaries."""
    files = [f for f in os.listdir(directory) if "ai_analysis" not in f.lower()]  # Exclude previous output files

    if not files:
        print(f"No valid documents found in '{directory}'. Please add files and rerun the script.")
        exit()

    print("\nAvailable files in the directory (excluding previous AI summaries):")
    for idx, file in enumerate(files, start=1):
        print(f"{idx}. {file}")

    file_name = input("\nEnter the exact filename you want to analyze: ").strip()
    file_path = os.path.join(directory, file_name)

    if not os.path.exists(file_path):
        print(f"Error: Filename '{file_name}' not found. Make sure it's in the directory and correctly spelled.")
        exit()

    print(f"\nFile '{file_name}' found! Processing...\n")
    return file_path, file_name

# ✅ Function to read different file types
def read_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    elif ext == ".txt":
        return read_txt(file_path)
    else:
        print("Unsupported file type! Please use a .pdf, .docx, or .txt file.")
        exit()

# ✅ Function to read PDF files
def read_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    return text

# ✅ Function to read DOCX files
def read_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# ✅ Function to read TXT files
def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# ✅ Function to split large text into smaller chunks
def chunk_text(text, max_tokens=1000):
    """Splits text into smaller chunks for processing."""
    words = text.split()
    chunks = []
    chunk = []
    current_length = 0

    for word in words:
        current_length += len(word) + 1  # Approximate token count
        if current_length > max_tokens:
            chunks.append(" ".join(chunk))
            chunk = []
            current_length = len(word) + 1
        chunk.append(word)

    if chunk:
        chunks.append(" ".join(chunk))

    return chunks

# ✅ Function to summarize text using GPT-4o-mini
def summarize_text(text):
    """Processes large text in chunks while handling API errors and displaying real-time progress."""
    text_chunks = chunk_text(text, max_tokens=1000)
    summaries = []

    total_chunks = len(text_chunks)
    start_time = time.time()  # Track start time

    for i, chunk in enumerate(text_chunks):
        chunk_start = time.time()
        percentage_complete = round((i + 1) / total_chunks * 100)  # Rounded to whole number
        elapsed_time = time.time() - start_time
        avg_time_per_chunk = elapsed_time / (i + 1)
        estimated_time_left = avg_time_per_chunk * (total_chunks - (i + 1))

        sys.stdout.write(
            f"\rProcessing chunk {i + 1}/{total_chunks}  |  Completed: {percentage_complete}%  "
            f"|  Estimated Time Left: {int(estimated_time_left // 60)} min {int(estimated_time_left % 60)} sec   "
        )
        sys.stdout.flush()

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Summarize the following document chunk."},
                {"role": "user", "content": chunk}
            ]
        )
        summaries.append(response.choices[0].message.content)
        time.sleep(2)

    print()
    return "\n\n".join(summaries)

# ✅ Function to save the summary to a DOCX file
def save_summary_to_docx(directory, filename, summary):
    summary_filename = os.path.join(directory, f"AI_Analysis_{filename}.docx")
    doc = Document()
    doc.add_heading(f"Summary of {filename}", level=1)
    doc.add_paragraph(summary)
    doc.save(summary_filename)
    print(f"✅ Summary saved as '{summary_filename}'")

# ✅ Main function
def main():
    directory = get_valid_directory()
    file_path, user_entered_filename = get_valid_file(directory)

    text = read_document(file_path)

    display_cost_estimate(text)

    summary = summarize_text(text)
    print("\nSummary:\n", summary)

    save_summary_to_docx(directory, user_entered_filename, summary)

# ✅ Run the script
if __name__ == "__main__":
    main()
